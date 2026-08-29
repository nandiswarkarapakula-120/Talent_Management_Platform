"""
Talent Management Platform - File Reading Utility
Extracts plain text from uploaded resume / profile PDFs, DOCX, or TXT files
so users can upload a file instead of typing/pasting content.
"""

import io


def extract_text_from_upload(uploaded_file):
    """
    Accepts a Streamlit UploadedFile object and returns extracted plain text.
    Supports .pdf, .docx, .txt
    Returns (text, error_message). If error_message is not None, text will be "".
    """
    if uploaded_file is None:
        return "", "No file provided."

    name = uploaded_file.name.lower()
    raw = uploaded_file.getvalue()

    try:
        if name.endswith(".pdf"):
            return _extract_pdf(raw), None
        elif name.endswith(".docx"):
            return _extract_docx(raw), None
        elif name.endswith(".txt"):
            return raw.decode("utf-8", errors="ignore"), None
        else:
            return "", "Unsupported file type. Please upload a PDF, DOCX, or TXT file."
    except Exception as e:
        return "", f"Could not read file: {e}"


def _extract_pdf(raw_bytes):
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(raw_bytes))
    text_parts = []
    for page in reader.pages:
        try:
            text_parts.append(page.extract_text() or "")
        except Exception:
            continue
    text = "\n".join(text_parts).strip()
    if not text:
        raise ValueError("No extractable text found (the PDF may be a scanned image).")
    return text


def _extract_docx(raw_bytes):
    import docx
    doc = docx.Document(io.BytesIO(raw_bytes))
    paragraphs = [p.text for p in doc.paragraphs]
    # Also pull table content (many resumes use tables for layout)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    text = "\n".join(paragraphs).strip()
    if not text:
        raise ValueError("No extractable text found in this DOCX file.")
    return text
